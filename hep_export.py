"""
hep_export.py - Export search results to PDF or Word document.
Called from hep_search.py /export route.
"""

import io
import re
from datetime import datetime


# ── Shared helpers ────────────────────────────────────────────────────────────

def _page_ref(c: dict) -> str:
    if c.get("page_start") and c.get("page_end"):
        return f"pp. {c['page_start']}–{c['page_end']}"
    if c.get("page_start"):
        return f"p. {c['page_start']}"
    return ""


def _clean_answer(text: str) -> str:
    """Strip markdown syntax for plain-text contexts."""
    text = re.sub(r'^#{1,3}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text.strip()


def _paragraphs(answer: str) -> list[str]:
    """Split answer into paragraphs, stripping markdown headers."""
    paras = []
    for block in answer.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        # Convert ## header to bold-style paragraph marker
        block = re.sub(r'^#{1,3}\s+', '', block)
        block = re.sub(r'\*\*(.+?)\*\*', r'\1', block)
        block = re.sub(r'\*(.+?)\*', r'\1', block)
        paras.append(block)
    return paras


# ── PDF export ────────────────────────────────────────────────────────────────

_PDF_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Source+Sans+Pro:ital,wght@0,300;0,400;0,600;1,400&display=swap');

* { box-sizing: border-box; margin: 0; padding: 0; }

body {
    font-family: 'Source Sans Pro', 'Helvetica Neue', Arial, sans-serif;
    font-size: 10.5pt;
    line-height: 1.65;
    color: #1a1a1a;
    background: #ffffff;
    padding: 0;
}

.page {
    max-width: 680px;
    margin: 0 auto;
    padding: 48px 48px 64px;
}

/* Header bar */
.doc-header {
    border-bottom: 3px solid #8B388D;
    padding-bottom: 16px;
    margin-bottom: 28px;
}

.org-name {
    font-size: 8pt;
    font-weight: 600;
    color: #8B388D;
    text-transform: uppercase;
    letter-spacing: 0.18em;
    margin-bottom: 4px;
}

.doc-title {
    font-size: 8pt;
    color: #666;
    letter-spacing: 0.05em;
}

/* Query */
.query-label {
    font-size: 7.5pt;
    font-weight: 600;
    color: #8B388D;
    text-transform: uppercase;
    letter-spacing: 0.15em;
    margin-bottom: 6px;
}

.query-text {
    font-size: 13pt;
    font-weight: 600;
    color: #1a1a1a;
    line-height: 1.4;
    margin-bottom: 28px;
    padding-bottom: 20px;
    border-bottom: 1px solid #e0d0e0;
}

/* Answer */
.answer-label {
    font-size: 7.5pt;
    font-weight: 600;
    color: #8B388D;
    text-transform: uppercase;
    letter-spacing: 0.15em;
    margin-bottom: 12px;
}

.answer-body p {
    margin-bottom: 11pt;
    font-size: 10.5pt;
    line-height: 1.65;
    color: #1a1a1a;
}

.answer-body p:last-child { margin-bottom: 0; }

sup {
    color: #8B388D;
    font-size: 7pt;
    font-weight: 600;
    vertical-align: super;
}

/* Contradictions */
.contra-section {
    margin-top: 24px;
    padding: 14px 16px;
    border: 1px solid #c87830;
    border-left: 4px solid #c87830;
    border-radius: 4px;
    background: #fffaf5;
}

.contra-label {
    font-size: 7.5pt;
    font-weight: 600;
    color: #c87830;
    text-transform: uppercase;
    letter-spacing: 0.15em;
    margin-bottom: 10px;
}

.contra-item { margin-bottom: 10px; }
.contra-item:last-child { margin-bottom: 0; }
.contra-topic { font-weight: 600; font-size: 9.5pt; color: #c87830; margin-bottom: 4px; }
.contra-side { font-size: 9pt; color: #333; padding-left: 10px; border-left: 2px solid #e8c090; margin-bottom: 3px; line-height: 1.5; }
.contra-num { color: #8B388D; font-weight: 600; }

/* Sources */
.sources-section {
    margin-top: 32px;
    padding-top: 20px;
    border-top: 2px solid #e0d0e0;
}

.sources-label {
    font-size: 7.5pt;
    font-weight: 600;
    color: #666;
    text-transform: uppercase;
    letter-spacing: 0.15em;
    margin-bottom: 14px;
}

.source-item {
    display: flex;
    gap: 10px;
    margin-bottom: 12px;
    padding-bottom: 12px;
    border-bottom: 1px solid #f0e8f0;
}

.source-item:last-child { border-bottom: none; margin-bottom: 0; }

.source-num {
    font-size: 8pt;
    font-weight: 700;
    color: #8B388D;
    min-width: 22px;
    padding-top: 1px;
}

.source-body { flex: 1; }

.source-title {
    font-size: 9.5pt;
    font-weight: 600;
    color: #1a1a1a;
    margin-bottom: 2px;
    line-height: 1.4;
}

.source-meta {
    font-size: 8.5pt;
    color: #666;
}

.source-excerpt {
    font-size: 8.5pt;
    color: #555;
    font-style: italic;
    margin-top: 4px;
    padding-left: 8px;
    border-left: 2px solid #d0b8d0;
    line-height: 1.5;
}

.evidence-tag {
    display: inline-block;
    font-size: 7pt;
    font-weight: 600;
    color: #8B388D;
    border: 1px solid #c0a0c0;
    border-radius: 3px;
    padding: 1px 5px;
    margin-left: 6px;
    vertical-align: middle;
    text-transform: uppercase;
    letter-spacing: 0.06em;
}

/* Footer */
.doc-footer {
    margin-top: 40px;
    padding-top: 12px;
    border-top: 1px solid #e0d0e0;
    font-size: 7.5pt;
    color: #999;
    display: flex;
    justify-content: space-between;
}
"""


def _answer_to_html_paras(answer: str) -> str:
    """Convert answer markdown to HTML paragraphs with superscript citations."""
    html_paras = []
    for block in answer.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        # Headers → bold paragraph
        block = re.sub(r'^#{1,3}\s+(.+)', r'<strong>\1</strong>', block)
        # Bold / italic
        block = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', block)
        block = re.sub(r'\*(.+?)\*', r'<em>\1</em>', block)
        # Citations [n] → superscript
        block = re.sub(r'\[(\d+)\]', r'<sup>[\1]</sup>', block)
        html_paras.append(f"<p>{block}</p>")
    return "\n".join(html_paras)


def export_pdf(query: str, answer: str, citations: list, contradictions: list) -> bytes:
    from weasyprint import HTML, CSS

    date_str = datetime.now().strftime("%-d %B %Y")

    # Contradiction HTML block
    contra_html = ""
    if contradictions:
        sides_html = ""
        for ct in contradictions:
            sides = "".join(
                f'<div class="contra-side"><span class="contra-num">[{s["index"]}]</span> {s["claim"]}</div>'
                for s in ct.get("sides", [])
            )
            sides_html += f'<div class="contra-item"><div class="contra-topic">{ct["topic"]}</div>{sides}</div>'
        contra_html = f'<div class="contra-section"><div class="contra-label">⚠ Conflicting evidence detected</div>{sides_html}</div>'

    # Citations HTML
    sources_html = ""
    for c in citations:
        strength     = c.get("evidence_strength", 1)
        badge        = f'<span class="evidence-tag">{strength} source{"s" if strength != 1 else ""}</span>' if strength else ""
        page_ref     = _page_ref(c)
        meta_parts   = [c.get("authors") or "", c.get("year") or ""]
        if page_ref:
            meta_parts.append(page_ref)
        meta_str     = " · ".join(p for p in meta_parts if p)
        excerpt      = c.get("excerpt", "")
        excerpt_html = f'<div class="source-excerpt">{excerpt[:250]}</div>' if excerpt else ""
        cat          = c.get("category", "")
        cat_html     = f' · <em>{cat}</em>' if cat else ""

        sources_html += f"""
        <div class="source-item">
          <div class="source-num">[{c['index']}]</div>
          <div class="source-body">
            <div class="source-title">{c.get('title', 'Unknown')}{badge}</div>
            <div class="source-meta">{meta_str}{cat_html}</div>
            {excerpt_html}
          </div>
        </div>"""

    answer_html = _answer_to_html_paras(answer)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>HEP Research Library Export</title></head>
<body>
<div class="page">
  <div class="doc-header">
    <div class="org-name">Hope Education Project · Research Library</div>
    <div class="doc-title">Research Export · {date_str}</div>
  </div>

  <div class="query-label">Research Question</div>
  <div class="query-text">{query}</div>

  <div class="answer-label">Answer</div>
  <div class="answer-body">{answer_html}</div>

  {contra_html}

  <div class="sources-section">
    <div class="sources-label">Sources · {len(citations)} documents</div>
    {sources_html}
  </div>

  <div class="doc-footer">
    <span>Hope Education Project · hopeeducationproject.org</span>
    <span>{date_str}</span>
  </div>
</div>
</body>
</html>"""

    return HTML(string=html).write_pdf(stylesheets=[CSS(string=_PDF_CSS)])


# ── Word export ───────────────────────────────────────────────────────────────

def export_docx(query: str, answer: str, citations: list, contradictions: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HEP_PURPLE = RGBColor(0x8B, 0x38, 0x8D)
    MID_GREY   = RGBColor(0x66, 0x66, 0x66)
    AMBER      = RGBColor(0xC8, 0x78, 0x30)
    BLACK      = RGBColor(0x1A, 0x1A, 0x1A)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def _add_label(text: str, colour: RGBColor = HEP_PURPLE):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.font.size  = Pt(7.5)
        run.font.bold  = True
        run.font.color.rgb = colour
        run.font.name  = "Calibri"
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def _add_rule(colour_hex: str = "8B388D", thickness: int = 18):
        """Add a coloured bottom border to the last paragraph."""
        p   = doc.paragraphs[-1]._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(thickness))
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), colour_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Header ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("HOPE EDUCATION PROJECT · RESEARCH LIBRARY")
    r.font.size  = Pt(8)
    r.font.bold  = True
    r.font.color.rgb = HEP_PURPLE
    r.font.name  = "Calibri"
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"Research Export · {datetime.now().strftime('%-d %B %Y')}")
    r2.font.size  = Pt(8)
    r2.font.color.rgb = MID_GREY
    r2.font.name  = "Calibri"
    p2.paragraph_format.space_after = Pt(6)
    _add_rule("8B388D", 24)

    # ── Query ─────────────────────────────────────────────────────────────────
    _add_label("Research Question")
    q = doc.add_paragraph()
    qr = q.add_run(query)
    qr.font.size  = Pt(13)
    qr.font.bold  = True
    qr.font.color.rgb = BLACK
    qr.font.name  = "Calibri"
    q.paragraph_format.space_after = Pt(16)
    _add_rule("D0C0D0", 6)

    # ── Answer ────────────────────────────────────────────────────────────────
    _add_label("Answer")
    for para_text in _paragraphs(answer):
        # Strip inline [n] citations - keep text readable
        clean = re.sub(r'\[(\d+)\]', r'[\1]', para_text)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        # Handle inline citation numbers - render them as superscript
        parts = re.split(r'(\[\d+\])', clean)
        for part in parts:
            m = re.match(r'\[(\d+)\]', part)
            if m:
                r = p.add_run(m.group(0))
                r.font.size       = Pt(7)
                r.font.superscript = True
                r.font.color.rgb  = HEP_PURPLE
                r.font.name       = "Calibri"
            else:
                r = p.add_run(part)
                r.font.size  = Pt(10.5)
                r.font.color.rgb = BLACK
                r.font.name  = "Calibri"

    # ── Contradictions ────────────────────────────────────────────────────────
    if contradictions:
        _add_label("Conflicting Evidence Detected", AMBER)
        for ct in contradictions:
            tp = doc.add_paragraph()
            tr = tp.add_run(ct.get("topic", ""))
            tr.font.bold  = True
            tr.font.size  = Pt(9.5)
            tr.font.color.rgb = AMBER
            tr.font.name  = "Calibri"
            tp.paragraph_format.space_after = Pt(2)

            for side in ct.get("sides", []):
                sp = doc.add_paragraph(style="List Bullet")
                nr = sp.add_run(f"[{side['index']}] ")
                nr.font.bold  = True
                nr.font.color.rgb = HEP_PURPLE
                nr.font.name  = "Calibri"
                nr.font.size  = Pt(9.5)
                cr = sp.add_run(side.get("claim", ""))
                cr.font.size  = Pt(9.5)
                cr.font.color.rgb = BLACK
                cr.font.name  = "Calibri"
                sp.paragraph_format.space_after = Pt(2)

    # ── Sources ───────────────────────────────────────────────────────────────
    _add_label("Sources", MID_GREY)
    _add_rule("D0C0D0", 12)

    for c in citations:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.space_before = Pt(8)

        # Number
        nr = p.add_run(f"[{c['index']}]  ")
        nr.font.bold  = True
        nr.font.color.rgb = HEP_PURPLE
        nr.font.size  = Pt(9)
        nr.font.name  = "Calibri"

        # Title
        tr = p.add_run(c.get("title", "Unknown"))
        tr.font.bold  = True
        tr.font.size  = Pt(9.5)
        tr.font.color.rgb = BLACK
        tr.font.name  = "Calibri"

        # Evidence strength
        strength = c.get("evidence_strength", 1)
        if strength:
            er = p.add_run(f"  [{strength} source{'s' if strength != 1 else ''}]")
            er.font.size  = Pt(7.5)
            er.font.color.rgb = HEP_PURPLE
            er.font.name  = "Calibri"

        # Meta line
        meta_parts = [c.get("authors") or "", c.get("year") or ""]
        page_ref   = _page_ref(c)
        if page_ref:
            meta_parts.append(page_ref)
        cat = c.get("category", "")
        if cat:
            meta_parts.append(cat)
        meta_str = " · ".join(p for p in meta_parts if p)

        mp = doc.add_paragraph()
        mr = mp.add_run(meta_str)
        mr.font.size  = Pt(8.5)
        mr.font.color.rgb = MID_GREY
        mr.font.name  = "Calibri"
        mp.paragraph_format.space_before = Pt(0)
        mp.paragraph_format.space_after  = Pt(2)

        # Excerpt
        excerpt = c.get("excerpt", "")
        if excerpt:
            ep = doc.add_paragraph()
            er2 = ep.add_run(excerpt[:250] + ("…" if len(excerpt) > 250 else ""))
            er2.font.size   = Pt(8.5)
            er2.font.italic = True
            er2.font.color.rgb = MID_GREY
            er2.font.name   = "Calibri"
            ep.paragraph_format.space_before = Pt(0)
            ep.paragraph_format.space_after  = Pt(2)
            ep.paragraph_format.left_indent  = Inches(0.2)

    # ── Footer ────────────────────────────────────────────────────────────────
    _add_rule("D0C0D0", 6)
    fp = doc.add_paragraph()
    fr = fp.add_run(f"Hope Education Project · hopeeducationproject.org · {datetime.now().strftime('%-d %B %Y')}")
    fr.font.size  = Pt(7.5)
    fr.font.color.rgb = MID_GREY
    fr.font.name  = "Calibri"
    fp.paragraph_format.space_before = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
